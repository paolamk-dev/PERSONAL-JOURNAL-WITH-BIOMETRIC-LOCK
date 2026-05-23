#!/usr/bin/env python3
"""
Visual Identity Guidelines Generator for MAKAMTE Smart Climate (MSC)
IoT-Based Smart Air Conditioning Regulation Project
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_cell_background(cell, fill_color):
    """Set background color for table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_cover_page(doc):
    """Create professional cover page"""
    # Add space at top
    for _ in range(8):
        doc.add_paragraph()

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Visual Identity Guidelines')
    run.font.size = Pt(42)
    run.font.color.rgb = RGBColor(0, 102, 153)  # Primary blue
    run.font.bold = True

    doc.add_paragraph()

    # Subtitle - Project name
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('MAKAMTE Smart Climate')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(41, 128, 185)
    run.font.bold = True

    doc.add_paragraph()

    # Project description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('IoT-Based Smart Air Conditioning Regulation')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    # Add space
    for _ in range(6):
        doc.add_paragraph()

    # Student info
    student = doc.add_paragraph()
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student.add_run('Prepared by:\nMAKAMTE TATSINKOU PAOLA')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date.add_run('May 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 120, 120)

    add_page_break(doc)

def create_table_of_contents(doc):
    """Create table of contents"""
    heading = doc.add_heading('Table of Contents', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    toc_items = [
        ('MAKAMTE Smart Climate Brand', '3'),
        ('Logo Concept & Imagery', '4'),
        ('Logo Configurations', '5'),
        ('    - Fixed Configuration', '5'),
        ('    - Flexible Configuration', '6'),
        ('Logo Structure & Clear Space', '7'),
        ('Design Elements', '8'),
        ('Primary Logo Variations', '9'),
        ('Color Palette', '10'),
        ('Color Meanings & Psychology', '11'),
        ('Typography Guidelines', '12'),
        ('Logo Usage Specifications', '13'),
        ('Logo Misusage', '14'),
        ('Application Examples', '15'),
        ('Conclusion', '16'),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        if item.startswith('    -'):
            p.paragraph_format.left_indent = Inches(0.5)
            item = item.strip()

        run1 = p.add_run(item)
        run1.font.size = Pt(12)

        # Add tab and page number
        run2 = p.add_run(f'\t{page}')
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(100, 100, 100)

    add_page_break(doc)

def create_brand_intro(doc):
    """Create brand introduction page"""
    heading = doc.add_heading('MAKAMTE Smart Climate', level=1)

    # Brand description
    p1 = doc.add_paragraph()
    run = p1.add_run(
        'MAKAMTE Smart Climate (MSC) represents the future of intelligent environmental '
        'control through IoT technology. Our brand embodies innovation, sustainability, '
        'and smart living.'
    )
    run.font.size = Pt(12)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    run = p2.add_run(
        'This Visual Identity Guidelines document serves as the comprehensive reference '
        'for maintaining brand consistency across all applications. It provides detailed '
        'specifications for logo usage, color schemes, typography, and design elements that '
        'define the MAKAMTE Smart Climate brand.'
    )
    run.font.size = Pt(12)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    run = p3.add_run(
        'The guidelines ensure that all visual communications—from digital interfaces to '
        'printed materials—maintain the professional quality and coherent identity that '
        'represents our commitment to excellence in IoT-based climate regulation.'
    )
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Brand values
    doc.add_heading('Brand Values', level=2)

    values = [
        ('Innovation', 'Pioneering IoT technology for intelligent climate control'),
        ('Sustainability', 'Energy-efficient solutions for environmental responsibility'),
        ('Intelligence', 'Smart algorithms that learn and adapt to user preferences'),
        ('Reliability', 'Dependable performance and consistent comfort'),
        ('Connectivity', 'Seamless integration with modern smart home ecosystems')
    ]

    for value, description in values:
        p = doc.add_paragraph(style='List Bullet')
        run1 = p.add_run(f'{value}: ')
        run1.bold = True
        run1.font.size = Pt(11)
        run2 = p.add_run(description)
        run2.font.size = Pt(11)

    add_page_break(doc)

def create_logo_imagery(doc):
    """Create logo imagery explanation page"""
    heading = doc.add_heading('Logo Concept & Imagery', level=1)

    doc.add_paragraph()

    # Logo concept
    doc.add_heading('The "M" Symbol', level=2)
    p1 = doc.add_paragraph()
    run = p1.add_run(
        'The MAKAMTE Smart Climate logo features a stylized "M" as its centerpiece, '
        'representing the brand name while incorporating visual elements that communicate '
        'our core mission.'
    )
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Design elements
    doc.add_heading('Design Elements', level=2)

    elements = [
        ('Flowing Lines', 'Represent air circulation and climate flow'),
        ('Geometric Precision', 'Symbolizes technological accuracy and IoT connectivity'),
        ('Cool Gradients', 'Evoke temperature regulation and cooling comfort'),
        ('Connected Structure', 'Reflects IoT network and smart home integration'),
        ('Modern Typography', 'Communicates innovation and forward-thinking design')
    ]

    for element, meaning in elements:
        p = doc.add_paragraph(style='List Bullet')
        run1 = p.add_run(f'{element}: ')
        run1.bold = True
        run1.font.size = Pt(11)
        run2 = p.add_run(meaning)
        run2.font.size = Pt(11)

    doc.add_paragraph()

    # Brand promise
    p2 = doc.add_paragraph()
    run = p2.add_run(
        'The overall design creates a sense of modern sophistication while maintaining '
        'approachability. The logo works equally well in digital applications (mobile apps, '
        'web interfaces) and physical materials (product labels, documentation).'
    )
    run.font.size = Pt(12)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    run = p3.add_run(
        'Consistency in logo application is paramount to building brand recognition and trust. '
        'All users must adhere to these guidelines to maintain the integrity of the MAKAMTE '
        'Smart Climate brand identity.'
    )
    run.font.size = Pt(12)
    run.italic = True

    add_page_break(doc)

def create_logo_configurations(doc):
    """Create logo configuration page"""
    heading = doc.add_heading('Logo Configurations', level=1)

    doc.add_paragraph()

    # Fixed Configuration
    doc.add_heading('Fixed Configuration', level=2)
    p1 = doc.add_paragraph()
    run = p1.add_run(
        'The MAKAMTE Smart Climate logo has two primary fixed configurations. When using '
        'fixed configurations, the logo mark and logotype are set in a fixed position relative '
        'to each other and cannot be altered.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    configs = [
        ('Horizontal Version', 'Primary configuration for most applications. Logo mark appears '
         'to the left of the wordmark. Ideal for headers, letterheads, and wide formats.'),
        ('Vertical/Stacked Version', 'Secondary configuration for space-constrained applications. '
         'Logo mark appears above the wordmark. Ideal for social media profiles, app icons, and '
         'square formats.')
    ]

    for config, desc in configs:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{config}\n')
        run1.bold = True
        run1.font.size = Pt(11)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    doc.add_paragraph()

    # Note box
    note = doc.add_paragraph()
    run = note.add_run('NOTE: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = note.add_run(
        'Both configurations maintain exact proportions between elements. Never attempt to '
        'recreate the logo by manually positioning elements.'
    )
    run2.font.size = Pt(10)
    run2.italic = True

    add_page_break(doc)

def create_flexible_configuration(doc):
    """Create flexible configuration page"""
    heading = doc.add_heading('Flexible Configuration', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'Flexible configurations allow the logo mark (the "M" symbol) to be used separately '
        'from the full wordmark in special circumstances. This configuration requires designer '
        'expertise and should be approved for use.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Use cases
    doc.add_heading('Appropriate Use Cases', level=2)

    uses = [
        'App icons and favicons where space is extremely limited',
        'Watermarks and background design elements',
        'Social media profile pictures',
        'Product physical branding (on hardware devices)',
        'Promotional materials where the full logo appears elsewhere'
    ]

    for use in uses:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(use)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Warning
    warning = doc.add_paragraph()
    run = warning.add_run('IMPORTANT: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)
    run2 = warning.add_run(
        'The standalone logo mark should never replace the full logo as the primary brand '
        'identifier. It serves as a supplementary design element only. When in doubt, use '
        'the full logo in fixed configuration.'
    )
    run2.font.size = Pt(11)

    add_page_break(doc)

def create_logo_structure(doc):
    """Create logo structure and clear space page"""
    heading = doc.add_heading('Logo Structure & Clear Space', level=1)

    doc.add_paragraph()

    # Clear space explanation
    doc.add_heading('Clear Space Requirements', level=2)
    p1 = doc.add_paragraph()
    run = p1.add_run(
        'To ensure maximum visibility and impact, the logo must be surrounded by adequate '
        'clear space. This protected area prevents other graphic elements, text, or images '
        'from crowding the logo and compromising its integrity.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Minimum clear space rules
    p2 = doc.add_paragraph()
    run1 = p2.add_run('Minimum Clear Space Rule:\n')
    run1.bold = True
    run1.font.size = Pt(11)
    run2 = p2.add_run(
        'The clear space around the logo should be equal to the height of the letter "M" in '
        'the logo mark. This measurement creates a consistent protection zone on all sides.'
    )
    run2.font.size = Pt(11)

    doc.add_paragraph()

    # Logo proportions
    doc.add_heading('Logo Proportions', level=2)
    p3 = doc.add_paragraph()
    run = p3.add_run(
        'The MAKAMTE Smart Climate logo is built on a precise grid system that ensures '
        'perfect balance and proportions. The relationship between the logo mark and '
        'wordmark follows these specifications:'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    proportions = [
        ('Horizontal Version', 'Logo mark width = 28% of total logo width; '
         'Wordmark width = 68% of total logo width; Space between = 4%'),
        ('Vertical Version', 'Logo mark height = 45% of total logo height; '
         'Wordmark height = 50% of total logo height; Space between = 5%')
    ]

    for prop, desc in proportions:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{prop}\n')
        run1.bold = True
        run1.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    doc.add_paragraph()

    # Scaling note
    note = doc.add_paragraph()
    run = note.add_run('SCALING NOTE: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = note.add_run(
        'Always scale the logo proportionally. Lock the aspect ratio when resizing. '
        'Never stretch or compress the logo horizontally or vertically.'
    )
    run2.font.size = Pt(10)
    run2.italic = True

    add_page_break(doc)

def create_design_elements(doc):
    """Create design elements overview page"""
    heading = doc.add_heading('Design Elements', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'The MAKAMTE Smart Climate brand includes multiple logo variations optimized for '
        'different use cases and backgrounds. Each variation maintains the core brand identity '
        'while providing flexibility for various applications.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Logo variations
    doc.add_heading('Available Logo Variations', level=2)

    variations = [
        ('Full Color', 'Primary version featuring complete color palette. Use on white or '
         'light backgrounds. Preferred for most applications.'),
        ('Reversed/White', 'Logo in white for dark backgrounds. Maintains brand recognition '
         'while providing necessary contrast.'),
        ('Monochrome Black', 'Single-color black version for specific printing needs or '
         'when color reproduction is not available.'),
        ('Monochrome White', 'Single-color white version for dark backgrounds when full '
         'color is not feasible.'),
        ('Grayscale', 'Grayscale version for black-and-white printing, photocopying, or '
         'fax transmission.')
    ]

    for variation, desc in variations:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{variation}\n')
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = RGBColor(0, 102, 153)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)
        doc.add_paragraph()

    # Selection guidance
    doc.add_heading('Selecting the Right Variation', level=2)
    p2 = doc.add_paragraph()
    run = p2.add_run(
        'Always choose the logo variation that provides maximum contrast and legibility '
        'against its background. The full-color version is preferred whenever possible. '
        'When applying logos to photographs or complex backgrounds, ensure sufficient '
        'contrast by using either the reversed white version on dark areas or the full-color '
        'version on light areas.'
    )
    run.font.size = Pt(11)

    add_page_break(doc)

def create_color_palette(doc):
    """Create color palette page"""
    heading = doc.add_heading('Color Palette', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'The MAKAMTE Smart Climate brand uses a carefully selected color palette that '
        'reflects our values of innovation, technology, and environmental consciousness. '
        'These colors must be reproduced accurately across all applications.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Color specifications
    doc.add_heading('Primary Brand Colors', level=2)

    # Create table for colors
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'

    colors = [
        {
            'name': 'Deep Ocean Blue',
            'hex': '#006699',
            'rgb': 'R:0 G:102 B:153',
            'cmyk': 'C:100 M:50 Y:0 K:20',
            'pantone': 'PMS 7691 C',
            'use': 'Primary brand color - Logo, headers, primary UI elements'
        },
        {
            'name': 'Sky Blue',
            'hex': '#2980B9',
            'rgb': 'R:41 G:128 B:185',
            'cmyk': 'C:78 M:45 Y:0 K:0',
            'pantone': 'PMS 7689 C',
            'use': 'Secondary color - Accents, highlights, interactive elements'
        },
        {
            'name': 'Cool Mint',
            'hex': '#16A085',
            'rgb': 'R:22 G:160 B:133',
            'cmyk': 'C:70 M:0 Y:50 K:0',
            'pantone': 'PMS 3272 C',
            'use': 'Tertiary color - Success states, environmental themes'
        },
        {
            'name': 'Tech Silver',
            'hex': '#95A5A6',
            'rgb': 'R:149 G:165 B:166',
            'cmyk': 'C:40 M:25 Y:25 K:0',
            'pantone': 'PMS Cool Gray 8 C',
            'use': 'Neutral color - Backgrounds, secondary text, dividers'
        },
        {
            'name': 'Charcoal',
            'hex': '#2C3E50',
            'rgb': 'R:44 G:62 B:80',
            'cmyk': 'C:80 M:60 Y:40 K:60',
            'pantone': 'PMS 432 C',
            'use': 'Text color - Primary typography, strong contrasts'
        }
    ]

    # Fill table with color information
    for idx, color in enumerate(colors):
        row = table.rows[idx]

        # Color name cell
        cell0 = row.cells[0]
        p = cell0.paragraphs[0]
        run = p.add_run(color['name'])
        run.bold = True
        run.font.size = Pt(11)

        # Color specs cell
        cell1 = row.cells[1]
        cell1.text = ''
        p = cell1.paragraphs[0]

        specs_text = f"HEX: {color['hex']}\n"
        specs_text += f"RGB: {color['rgb']}\n"
        specs_text += f"CMYK: {color['cmyk']}\n"
        specs_text += f"Pantone: {color['pantone']}\n"
        specs_text += f"Use: {color['use']}"

        run = p.add_run(specs_text)
        run.font.size = Pt(9)

    doc.add_paragraph()

    # Color accuracy note
    note = doc.add_paragraph()
    run = note.add_run('COLOR ACCURACY: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = note.add_run(
        'Efforts must be made to maintain exact color reproduction as specified above. '
        'Use Pantone colors for professional printing, CMYK for standard printing, RGB for '
        'digital displays, and HEX for web applications. Never deviate from these specifications.'
    )
    run2.font.size = Pt(10)
    run2.italic = True

    add_page_break(doc)

def create_color_meanings(doc):
    """Create color meanings page"""
    heading = doc.add_heading('Color Meanings & Psychology', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'Each color in the MAKAMTE Smart Climate palette has been chosen for its psychological '
        'impact and symbolic meaning within the context of IoT technology and climate control.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Color meanings
    meanings = [
        {
            'color': 'Deep Ocean Blue (#006699)',
            'meaning': 'Trust, Professionalism & Reliability',
            'description': 'Represents the dependability of our technology and the trust users '
                         'place in our systems. Blue is universally associated with technology, '
                         'stability, and cooling—perfectly aligned with air conditioning.'
        },
        {
            'color': 'Sky Blue (#2980B9)',
            'meaning': 'Innovation & Clarity',
            'description': 'Symbolizes clear skies, fresh air, and the transparent nature of our '
                         'IoT systems. This lighter blue communicates approachability and modern '
                         'innovation, making technology accessible to all users.'
        },
        {
            'color': 'Cool Mint (#16A085)',
            'meaning': 'Freshness & Sustainability',
            'description': 'Evokes the feeling of cool, fresh air and environmental consciousness. '
                         'This color represents our commitment to energy efficiency and sustainable '
                         'climate solutions that benefit both users and the planet.'
        },
        {
            'color': 'Tech Silver (#95A5A6)',
            'meaning': 'Modernity & Balance',
            'description': 'Represents the sleek, modern hardware of IoT devices and the balanced, '
                         'neutral approach our system takes. Silver suggests technological sophistication '
                         'without being overwhelming.'
        },
        {
            'color': 'Charcoal (#2C3E50)',
            'meaning': 'Strength & Precision',
            'description': 'Provides strong contrast and readability while suggesting the robust, '
                         'precise nature of our control algorithms. This dark tone grounds the lighter '
                         'colors and adds professional gravitas.'
        }
    ]

    for item in meanings:
        # Color name as subheading
        p = doc.add_paragraph()
        run = p.add_run(item['color'])
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 102, 153)

        # Meaning
        p2 = doc.add_paragraph()
        run = p2.add_run(f"{item['meaning']}\n")
        run.bold = True
        run.font.size = Pt(11)
        run2 = p2.add_run(item['description'])
        run2.font.size = Pt(10)

        doc.add_paragraph()

    # Combined impact
    doc.add_heading('Combined Visual Impact', level=2)
    p_combined = doc.add_paragraph()
    run = p_combined.add_run(
        'Together, these colors create a cohesive palette that communicates technological '
        'sophistication, environmental responsibility, and user-friendly innovation. The cool '
        'tones inherently suggest climate control and comfort, while the professional blues '
        'establish credibility in the IoT marketplace.'
    )
    run.font.size = Pt(11)

    add_page_break(doc)

def create_typography(doc):
    """Create typography guidelines page"""
    heading = doc.add_heading('Typography Guidelines', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'Typography is a critical component of the MAKAMTE Smart Climate brand identity. '
        'Consistent use of specified typefaces ensures professional communication and '
        'reinforces brand recognition.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Primary typeface
    doc.add_heading('Primary Typeface: Montserrat', level=2)

    p2 = doc.add_paragraph()
    run1 = p2.add_run('Montserrat Bold\n')
    run1.bold = True
    run1.font.size = Pt(14)
    run2 = p2.add_run('Used for: ')
    run2.font.size = Pt(10)
    run2.italic = True
    run3 = p2.add_run('Logo wordmark, main headings, call-to-action buttons')
    run3.font.size = Pt(10)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    run1 = p3.add_run('Montserrat Medium\n')
    run1.bold = True
    run1.font.size = Pt(12)
    run2 = p3.add_run('Used for: ')
    run2.font.size = Pt(10)
    run2.italic = True
    run3 = p3.add_run('Subheadings, section titles, emphasis text')
    run3.font.size = Pt(10)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    run1 = p4.add_run('Montserrat Regular\n')
    run1.font.size = Pt(11)
    run2 = p4.add_run('Used for: ')
    run2.font.size = Pt(10)
    run2.italic = True
    run3 = p4.add_run('Body text, descriptions, general content')
    run3.font.size = Pt(10)

    doc.add_paragraph()

    # Why Montserrat
    info = doc.add_paragraph()
    run = info.add_run('Why Montserrat? ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = info.add_run(
        'This geometric sans-serif typeface conveys modernity and clarity. Its clean lines '
        'and excellent readability make it perfect for both digital interfaces and print materials. '
        'Montserrat is freely available via Google Fonts, ensuring accessibility across all platforms.'
    )
    run2.font.size = Pt(10)

    doc.add_paragraph()

    # Secondary typeface
    doc.add_heading('Secondary Typeface: Open Sans', level=2)

    p5 = doc.add_paragraph()
    run = p5.add_run('Open Sans')
    run.font.size = Pt(11)
    run2 = p5.add_run('\nUsed for: ')
    run2.font.size = Pt(10)
    run2.italic = True
    run3 = p5.add_run('Secondary body text, technical specifications, fine print, data displays')
    run3.font.size = Pt(10)

    doc.add_paragraph()

    info2 = doc.add_paragraph()
    run = info2.add_run('Why Open Sans? ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = info2.add_run(
        'Open Sans provides excellent legibility at small sizes and complements Montserrat perfectly. '
        'Use it for longer text passages and technical documentation where readability is paramount.'
    )
    run2.font.size = Pt(10)

    doc.add_paragraph()

    # Fallback fonts
    doc.add_heading('System Fallback Fonts', level=2)

    p6 = doc.add_paragraph()
    run = p6.add_run(
        'If Montserrat or Open Sans are unavailable, use these system alternatives:\n'
        '• Windows: Segoe UI or Arial\n'
        '• macOS: San Francisco or Helvetica Neue\n'
        '• Web: System font stack (system-ui, -apple-system, BlinkMacSystemFont)'
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Typography rules
    doc.add_heading('Typography Rules', level=2)

    rules = [
        'Never use decorative or script fonts in official brand materials',
        'Maintain consistent hierarchy: Headings larger and bolder than body text',
        'Use sufficient line spacing (1.5x for body text, 1.2x for headings)',
        'Never stretch, condense, or artificially italicize the typefaces',
        'Ensure adequate contrast between text and background (minimum 4.5:1 ratio)'
    ]

    for rule in rules:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(rule)
        run.font.size = Pt(10)

    add_page_break(doc)

def create_logo_usage(doc):
    """Create logo usage specifications page"""
    heading = doc.add_heading('Logo Usage Specifications', level=1)

    doc.add_paragraph()

    # Minimum size requirements
    doc.add_heading('Minimum Size Requirements', level=2)

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'To maintain legibility and brand impact, the logo must never be reproduced smaller '
        'than the specified minimum sizes. These measurements ensure that all logo elements '
        'remain clear and readable.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Size specifications
    size_specs = [
        ('Horizontal Logo - Print', 'Minimum width: 1.5 inches (38mm)',
         'Ensures the wordmark remains legible on printed materials'),
        ('Horizontal Logo - Digital', 'Minimum width: 150 pixels',
         'Maintains clarity on screens and digital displays'),
        ('Vertical Logo - Print', 'Minimum width: 1.0 inch (25mm)',
         'Compact format remains readable when space is limited'),
        ('Vertical Logo - Digital', 'Minimum width: 100 pixels',
         'Appropriate for social media profiles and app icons'),
        ('Standalone Mark - Digital', 'Minimum size: 32×32 pixels',
         'Smallest acceptable size for favicons and small icons')
    ]

    for spec, size, note in size_specs:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{spec}\n')
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = RGBColor(0, 102, 153)
        run2 = p.add_run(f'{size}\n')
        run2.font.size = Pt(10)
        run2.bold = True
        run3 = p.add_run(note)
        run3.font.size = Pt(9)
        run3.italic = True
        doc.add_paragraph()

    # Background usage
    doc.add_heading('Background Usage Guidelines', level=2)

    p2 = doc.add_paragraph()
    run = p2.add_run('Full-Color Logo: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p2.add_run('Use on white or very light backgrounds only. Ensure sufficient contrast.')
    run2.font.size = Pt(10)

    p3 = doc.add_paragraph()
    run = p3.add_run('Reversed/White Logo: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p3.add_run('Use on dark backgrounds (Charcoal #2C3E50 or darker). Minimum background '
                      'darkness: 70% gray or darker.')
    run2.font.size = Pt(10)

    p4 = doc.add_paragraph()
    run = p4.add_run('Photos & Complex Backgrounds: ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p4.add_run('Apply logos to solid color areas within images. If necessary, add a '
                      'subtle backdrop (white box with 80% opacity for color logo, dark box '
                      'for white logo) to ensure visibility.')
    run2.font.size = Pt(10)

    doc.add_paragraph()

    # File formats
    doc.add_heading('Preferred File Formats', level=2)

    formats = [
        ('Print Applications', 'Vector formats: EPS, PDF, or AI files with embedded fonts'),
        ('Digital/Screen', 'PNG with transparent background (24-bit) or SVG for web'),
        ('Microsoft Office', 'High-resolution PNG (300 DPI) or EMF vector format'),
        ('Web Development', 'SVG (scalable) or PNG at 2x resolution for retina displays')
    ]

    for use_case, format_spec in formats:
        p = doc.add_paragraph(style='List Bullet')
        run1 = p.add_run(f'{use_case}: ')
        run1.bold = True
        run1.font.size = Pt(10)
        run2 = p.add_run(format_spec)
        run2.font.size = Pt(10)

    add_page_break(doc)

def create_logo_misusage(doc):
    """Create logo misusage page"""
    heading = doc.add_heading('Logo Misusage', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'To protect the integrity of the MAKAMTE Smart Climate brand, the following practices '
        'are strictly prohibited. Any deviation from the approved logo usage guidelines weakens '
        'brand recognition and professionalism.'
    )
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph()

    # Prohibited actions
    doc.add_heading('NEVER Do the Following:', level=2)

    prohibited = [
        ('DO NOT Stretch or Distort',
         'Never stretch, compress, or skew the logo. Always maintain the original aspect ratio. '
         'Lock proportions when resizing.'),

        ('DO NOT Change Colors',
         'Never alter the specified brand colors or substitute with different colors. Do not apply '
         'gradients, patterns, or filters to logo elements (except approved variations).'),

        ('DO NOT Rotate or Tilt',
         'The logo must always appear horizontal and level. Never rotate, tilt, or apply perspective '
         'effects.'),

        ('DO NOT Rearrange Elements',
         'Never change the position of logo elements relative to each other. Do not separate the '
         'logo mark from the wordmark in fixed configurations.'),

        ('DO NOT Add Effects',
         'Never add drop shadows, glows, bevels, embossing, or other visual effects to the logo. '
         'Use only the approved logo files.'),

        ('DO NOT Outline',
         'Never place borders, outlines, or frames around the logo. The logo should exist within '
         'its clear space without enclosure.'),

        ('DO NOT Recreate',
         'Never attempt to redraw, recreate, or typeset the logo. Always use official logo files. '
         'Do not substitute fonts.'),

        ('DO NOT Use Low Quality',
         'Never use low-resolution, pixelated, or blurry logo files. Always source logos from the '
         'official brand asset library.'),

        ('DO NOT Overlap Busy Areas',
         'Never place the logo over busy photographs or complex backgrounds without ensuring '
         'adequate contrast and legibility.'),

        ('DO NOT Reduce Below Minimum',
         'Never reproduce the logo smaller than the specified minimum sizes. Legibility must be '
         'maintained at all times.'),

        ('DO NOT Animate Incorrectly',
         'If animating the logo for digital use, maintain brand integrity. Avoid spinning, bouncing, '
         'or other unprofessional animation effects.'),

        ('DO NOT Alter Typography',
         'Never change the font, letter spacing, or styling of the wordmark. Use only official logo files.')
    ]

    for title, description in prohibited:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{title}\n')
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = RGBColor(192, 0, 0)
        run2 = p.add_run(description)
        run2.font.size = Pt(10)
        doc.add_paragraph()

    # Final warning
    warning = doc.add_paragraph()
    run = warning.add_run('IMPORTANT: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(192, 0, 0)
    run2 = warning.add_run(
        'When in doubt, consult these guidelines or contact the brand administrator. Protecting '
        'our visual identity is everyone\'s responsibility. Consistent, correct logo usage builds '
        'trust and recognition in the marketplace.'
    )
    run2.font.size = Pt(11)

    add_page_break(doc)

def create_application_examples(doc):
    """Create application examples page"""
    heading = doc.add_heading('Application Examples', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'The MAKAMTE Smart Climate visual identity extends across numerous touchpoints. '
        'This section provides guidance on applying the brand to common materials and platforms.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Digital applications
    doc.add_heading('Digital Applications', level=2)

    digital_apps = [
        ('Mobile Application',
         'Use horizontal logo in app header. Standalone mark for app icon. Ensure minimum 150px width. '
         'Apply on white background with primary blue (#006699) navigation elements.'),

        ('Website',
         'Horizontal logo in top-left header position. SVG format for crisp scaling. Link logo to '
         'homepage. Minimum 180px width. White background preferred.'),

        ('Social Media Profiles',
         'Vertical/stacked logo for square profile pictures (Facebook, LinkedIn, Instagram). '
         'Standalone mark for small icons (Twitter/X favicon). Export at 2x resolution (500×500px minimum).'),

        ('Email Signatures',
         'Horizontal logo, 150-200px width. PNG format. Link to company website. Place at '
         'signature beginning or end on white background.'),

        ('Digital Presentations',
         'Horizontal logo in slide master header or footer. Consistent placement across all slides. '
         'Vertical logo for title slides. Minimum 1.5 inches wide.')
    ]

    for app, guidance in digital_apps:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{app}\n')
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = RGBColor(0, 102, 153)
        run2 = p.add_run(guidance)
        run2.font.size = Pt(9)
        doc.add_paragraph()

    # Print applications
    doc.add_heading('Print Applications', level=2)

    print_apps = [
        ('Business Cards',
         'Horizontal or vertical logo depending on card orientation. Minimum 1 inch width. '
         'Vector EPS or high-res PDF. Full color on white card stock.'),

        ('Letterhead',
         'Horizontal logo in top-left or centered header. 1.5-2 inches wide. CMYK color mode. '
         'Consistent footer placement optional.'),

        ('Product Packaging',
         'Prominent logo placement on primary display panel. Minimum 2 inches for boxes. '
         'Consider reversed/white logo for dark packaging. Use Pantone colors for color accuracy.'),

        ('Technical Documentation',
         'Logo in document header, 1.5 inches wide. Repeated on each page or first page only. '
         'Pair with brand colors for section headers.'),

        ('Promotional Materials',
         'Logo sizing depends on material size. Maintain 1:8 ratio (logo width to material width). '
         'Use full-color version prominently.')
    ]

    for app, guidance in print_apps:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{app}\n')
        run1.bold = True
        run1.font.size = Pt(11)
        run1.font.color.rgb = RGBColor(0, 102, 153)
        run2 = p.add_run(guidance)
        run2.font.size = Pt(9)
        doc.add_paragraph()

    # Physical product applications
    doc.add_heading('Physical Product Applications', level=2)

    p2 = doc.add_paragraph()
    run = p2.add_run(
        'For IoT hardware devices, sensors, and thermostats, the standalone "M" logo mark may be '
        'used on physical products where space is limited. Embossing, etching, or printing in white '
        'on dark device surfaces is acceptable. The full logo should appear on device packaging and '
        'accompanying materials.'
    )
    run.font.size = Pt(10)

    add_page_break(doc)

def create_conclusion(doc):
    """Create conclusion page"""
    heading = doc.add_heading('Conclusion', level=1)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    run = p1.add_run(
        'These Visual Identity Guidelines represent the foundation of the MAKAMTE Smart Climate '
        'brand. Consistent adherence to these standards ensures that our brand communicates '
        'professionalism, innovation, and reliability across all touchpoints.'
    )
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Key takeaways
    doc.add_heading('Key Takeaways', level=2)

    takeaways = [
        'Always use official logo files—never recreate or modify the logo',
        'Respect minimum size requirements to maintain legibility',
        'Use the correct logo variation for each background and application',
        'Maintain exact brand colors using provided specifications',
        'Apply approved typography consistently across all materials',
        'Protect the logo with adequate clear space',
        'When uncertain, refer to these guidelines or seek approval'
    ]

    for takeaway in takeaways:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(takeaway)
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Brand stewardship
    doc.add_heading('Brand Stewardship', level=2)

    p2 = doc.add_paragraph()
    run = p2.add_run(
        'Every team member, partner, and stakeholder who represents MAKAMTE Smart Climate is a '
        'brand steward. Your commitment to these guidelines helps build a strong, recognizable, '
        'and trustworthy brand in the competitive IoT marketplace.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    run = p3.add_run(
        'Consistent visual identity creates familiarity and trust with our users. It distinguishes '
        'our innovative IoT climate solutions from competitors and reinforces the quality and '
        'professionalism that MAKAMTE Smart Climate represents.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Future updates
    doc.add_heading('Living Document', level=2)

    p4 = doc.add_paragraph()
    run = p4.add_run(
        'These guidelines may be updated periodically to reflect brand evolution and new applications. '
        'Always refer to the latest version of this document for current specifications. Future updates '
        'will maintain core brand elements while potentially adding new variations or use cases as the '
        'MAKAMTE Smart Climate ecosystem grows.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Contact information
    doc.add_heading('Questions or Clarifications', level=2)

    p5 = doc.add_paragraph()
    run = p5.add_run(
        'For questions about logo usage, requests for logo files, or clarification on any aspect '
        'of the visual identity guidelines, please contact:'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    contact = doc.add_paragraph()
    run = contact.add_run('MAKAMTE TATSINKOU PAOLA\n')
    run.bold = True
    run.font.size = Pt(11)
    run2 = contact.add_run('Brand & Visual Identity Manager\n')
    run2.font.size = Pt(10)
    run3 = contact.add_run('MAKAMTE Smart Climate (MSC) Project\n')
    run3.font.size = Pt(10)
    run4 = contact.add_run('Software Engineering Department\n')
    run4.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Closing statement
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run('Thank you for helping maintain the integrity\nof the MAKAMTE Smart Climate brand.')
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()

    # Document info
    doc_info = doc.add_paragraph()
    doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_info.add_run('—\nVisual Identity Guidelines v1.0\nMay 2026')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

def main():
    """Main function to create the Visual Identity Guidelines document"""
    print("Creating MAKAMTE Smart Climate Visual Identity Guidelines...")

    # Create document
    doc = Document()

    # Set document properties
    doc.core_properties.title = 'MAKAMTE Smart Climate - Visual Identity Guidelines'
    doc.core_properties.author = 'MAKAMTE TATSINKOU PAOLA'
    doc.core_properties.subject = 'Brand Visual Identity Guidelines'
    doc.core_properties.keywords = 'MAKAMTE, Smart Climate, MSC, IoT, Branding, Visual Identity, Logo'

    # Create all sections
    print("  ✓ Creating cover page...")
    create_cover_page(doc)

    print("  ✓ Creating table of contents...")
    create_table_of_contents(doc)

    print("  ✓ Creating brand introduction...")
    create_brand_intro(doc)

    print("  ✓ Creating logo imagery section...")
    create_logo_imagery(doc)

    print("  ✓ Creating logo configurations...")
    create_logo_configurations(doc)

    print("  ✓ Creating flexible configuration...")
    create_flexible_configuration(doc)

    print("  ✓ Creating logo structure...")
    create_logo_structure(doc)

    print("  ✓ Creating design elements...")
    create_design_elements(doc)

    print("  ✓ Creating color palette...")
    create_color_palette(doc)

    print("  ✓ Creating color meanings...")
    create_color_meanings(doc)

    print("  ✓ Creating typography guidelines...")
    create_typography(doc)

    print("  ✓ Creating logo usage specifications...")
    create_logo_usage(doc)

    print("  ✓ Creating logo misusage guidelines...")
    create_logo_misusage(doc)

    print("  ✓ Creating application examples...")
    create_application_examples(doc)

    print("  ✓ Creating conclusion...")
    create_conclusion(doc)

    # Save document
    output_path = '/home/daytona/project/MAKAMTE_Smart_Climate_Visual_Identity_Guidelines.docx'
    doc.save(output_path)

    print(f"\n✅ Document created successfully!")
    print(f"📄 Saved to: {output_path}")
    print(f"\n📊 Document Statistics:")
    print(f"   - Total pages: ~16")
    print(f"   - Format: Fully editable .docx")
    print(f"   - Ready for: Printing and submission")

    return output_path

if __name__ == '__main__':
    main()
